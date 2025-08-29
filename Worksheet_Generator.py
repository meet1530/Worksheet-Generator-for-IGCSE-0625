# -*- coding: utf-8 -*-
"""
Created on Thu Mar 06 10:55:56 2025

@author: meetp
"""

import os
import google.generativeai as genai
from docx import Document
import re
from datetime import datetime

# Configure Gemini API (Use environment variable for security)
genai.configure(api_key="AIzaSyD9UiWN3O-6meRw_CMceqiNlJWExmzbTsw")

def load_syllabus_from_docx(file_path):
    syllabus = {}
    doc = Document(file_path)
    current_topic = None
    
    for para in doc.paragraphs:
        line = para.text.strip()
        match = re.match(r"^(\d+\.\d+)\s+(.*)$", line)  # Detect main topics (e.g., "1.1 Physical Quantities")
        if match:
            current_topic, topic_name = match.groups()
            syllabus[current_topic] = {"name": topic_name, "subtopics": []}
        elif current_topic and line:  # If it's indented or follows a topic, treat as subtopic
            syllabus[current_topic]["subtopics"].append(line)
    
    return syllabus

# Load syllabus from Word file
SYLLABUS_FILE = "0625_Syllabus.docx"  # Ensure this file is in the same directory
SYLLABUS = load_syllabus_from_docx(SYLLABUS_FILE)

def get_questions_and_answers(topic_number):
    if topic_number not in SYLLABUS:
        raise ValueError(f"Invalid topic number: {topic_number}. Available topics: {list(SYLLABUS.keys())}")
    
    topic_info = SYLLABUS[topic_number]
    topic_name = topic_info["name"]
    subtopics = topic_info["subtopics"]
    
    prompt = (
        f"Generate a structured IBDP Physics HL worksheet for topic: {topic_name}. "
        "Ensure the worksheet includes exactly: "
        "1. Forteen multiple-choice questions (MCQs) with four options each, "
        # "2. Five short/long answer questions, "
        # "3. Two unique, story-based physics puzzles. "
        "Each section should start with the following heading (Do not make the headings bold): "
        "MULTIPLE-CHOICE QUESTIONS (MCQs) "
        "Keep the difficulty very hard and keep numericals and conceptual questions"
        # "SHORT/LONG ANSWER QUESTIONS "
        # "STORY-BASED PHYSICS PUZZLES "
        # "Ensure that each question has a marking scheme, with marks indicated at the end of each question in parentheses, e.g., (2 marks). "
        # "For MCQs, ensure the options are clearly listed one below the other, and the marks appear immediately at the end of the question before the options. "
        # "Do NOT provide an answer key in this response. Generate only questions."
    )

    
    if subtopics:
        prompt += f" The questions should cover these subtopics: {', '.join(subtopics)}."
    
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content([prompt])
    questions = response.text.split("\n") if response and response.text else ["No questions generated."]
    
    answer_prompt = f"Provide the answer key for the following IGCSE Physics questions:\n" + "\n".join(questions) + "Keep formating of all the text same (Do not make anything bold)"
    answer_response = model.generate_content([answer_prompt])
    answers = answer_response.text.split("\n") if answer_response and answer_response.text else ["No answers generated."]
    
    return [q.strip() for q in questions if q.strip()], [a.strip() for a in answers if a.strip()]

def generate_worksheet_docx(topic_number, topic_name, questions):
    date_today = datetime.today().strftime('%Y%m%d_%H%M%S')
    filename = f"worksheet_{date_today}_topic_{topic_number}.docx"
    
    doc = Document()
    doc.add_heading(f"Worksheet - {datetime.today().strftime('%Y-%m-%d %H:%M:%S')}", level=1)
    doc.add_heading(f"{topic_number}: {topic_name}", level=2)
    
    sections = ["MULTIPLE-CHOICE QUESTIONS (MCQs)", "SHORT/LONG ANSWER QUESTIONS", "STORY-BASED PHYSICS PUZZLES"]
    
    for q in questions:
        if q in sections:
            doc.add_heading(q, level=3)
        else:
            doc.add_paragraph(q)
    
    doc.save(filename)
    return filename

def generate_answer_key_docx(topic_number, answers):
    date_today = datetime.today().strftime('%Y%m%d_%H%M%S')
    filename = f"answer_key_{date_today}_topic_{topic_number}.docx"
    
    doc = Document()
    doc.add_heading("Answer Key", level=1)
    
    for ans in answers:
        doc.add_paragraph(ans)
    
    doc.save(filename)
    return filename

def generate_worksheet(topic_number):
    topic_info = SYLLABUS.get(topic_number, {"name": "Unknown Topic", "subtopics": []})
    topic_name = topic_info["name"]
    questions, answers = get_questions_and_answers(topic_number)
    
    worksheet_file = generate_worksheet_docx(topic_number, topic_name, questions)
    answer_key_file = generate_answer_key_docx(topic_number, answers)
    
    return worksheet_file, answer_key_file

# Example usage
topic_number = "7.1"
generate_worksheet(topic_number)
